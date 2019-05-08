from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt
import requests
import urllib
import urllib.request
from bs4 import BeautifulSoup
import string
from WordScrape import getword



month = ""
str(month)
day = ""
str(day)
month = "june" #Wrap GUI Control
day = "3"


url = "https://www.timeanddate.com/on-this-day" + "/" + month + "/" + day

request = requests.get(url)		#<---------- Move this into a seperate function
html_content = request.text

soup = BeautifulSoup(html_content, "html.parser")

#Find the tag and its attribute
x = soup.find("h3", class_ ='otd-ttl')
y = x.find_next("h3", class_ ='otd-ttl')
g = y.find_next("h3", class_='otd-ttl')
h = g.find_next("h3", class_ = 'otd-ttl')
e = h.find_next("h3", class_ = 'otd-ttl')


def parsetotext(foundsoup):
	k = []
	for child in foundsoup:
		k.append(child)
	l = k[1]
	i = l.strip(' ')
	rcb = []
	for zx in k[0]:
		rcb.append(zx)
	return rcb[0]+"-"+i
#------------------------------------------------------#
#

#8:14 - YESSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS
                                     

#Open Document
document = Document()
section = document.sections[0]

#FORMATTING - LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENTATION.LANDSCAPE 
section.page_width = new_width
section.page_height = new_height

#Title styling
style1 = document.styles['Normal']
font1 = style1.font
font1.name = 'Times New Roman'
font1.size = Pt(30)

#Styling
header = section.header
paragraphtitle = header.paragraphs[0]
paragraphtitle.text = "The Verona Times"
paragraphtitle.style = document.styles['Normal']
paragraphtitle.paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


#------------------STYLES------------------------------------------#

#Add separate stylings for content types

#HEADING
obj_styles = document.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(12)
obj_font.name = 'Century Gothic'

#CONTENT
obj_charstyle2 = obj_styles.add_style('Contentstyle', WD_STYLE_TYPE.CHARACTER)
obj_font2 = obj_charstyle2.font
obj_font2.size = Pt(12)
obj_font2.name = 'Century Gothic'

#TITLE


#Table Formats
table = document.add_table(rows=5, cols=5)


Mdate = " 11/5"
Tdate = " 11/6"
WDate = " 11/7"
ThDate = " 11/8"
FDate = " 11/9"

Mevent = parsetotext(x)
Tevent = parsetotext(y)
Wevent = parsetotext(g)
Thevent = parsetotext(h)
Fevent = parsetotext(e)



 #---------------------------------------------------------------------------------------#
#Monday
Monday = table.cell(0,0)
Mondayword = table.cell(3,0)
WordDescription = table.cell(2,0)
Mondayevent = table.cell(1,0)
# Bolding
Monday.paragraphs[0].add_run("Monday"+str(Mdate), style = 'Contentstyle').bold = True
Mondayword.paragraphs[0].add_run('Verminllion(n) - A vivid red to reddish orange color', style = 'CommentsStyle')
WordDescription.paragraphs[0].add_run("Word of the Day: " + getword(2019,5,6), style = 'CommentsStyle')
Mondayevent.paragraphs[0].add_run(Mevent, style = 'CommentsStyle')



 #---------------------------------------------------------------------------------------#
Tuesday = table.cell(0,1)
Tuesdayword = table.cell(3, 1)
WordDescription2 = table.cell(2,1)
Tuesdayevent = table.cell(1,1)
Tuesday.paragraphs[0].add_run("Tuesday"+str(Tdate), style = 'Contentstyle').bold = True
WordDescription2.paragraphs[0].add_run("Word of the Day: "+ getword(2019,5,7), style = 'CommentsStyle')
Tuesdayword.paragraphs[0].add_run('Verminllion(n) - A vivid red to reddish orange color', style = 'CommentsStyle')
Tuesdayevent.paragraphs[0].add_run(Tevent, style = 'CommentsStyle')


 #---------------------------------------------------------------------------------------#
Wednesday = table.cell(0,2)
Wednesdayword = table.cell(3, 2)
WordDescription3 = table.cell(2,2)
Wednesdayevent = table.cell(1,2)
Wednesday.paragraphs[0].add_run("Wednesday"+str(WDate), style = 'Contentstyle').bold = True
WordDescription3.paragraphs[0].add_run("Word of the Day: " + getword(2019,5,8), style = 'CommentsStyle')
Wednesdayword.paragraphs[0].add_run('Verminllion(n) - A vivid red to reddish orange color', style = 'CommentsStyle')
Wednesdayevent.paragraphs[0].add_run(Wevent, style = 'CommentsStyle')

 
 #---------------------------------------------------------------------------------------#
Thursday= table.cell(0,3)
Thursdayword = table.cell(3, 3)
WordDescription4 = table.cell(2,3)
Thursdayevent = table.cell(1,3)
Thursday.paragraphs[0].add_run("Thursday"+str(ThDate), style = 'Contentstyle').bold = True
WordDescription4.paragraphs[0].add_run("Word of the Day: " + getword(2019,5,9), style = 'CommentsStyle')
Thursdayword.paragraphs[0].add_run('Verminllion(n) - A vivid red to reddish orange color', style = 'CommentsStyle')
Thursdayevent.paragraphs[0].add_run(Thevent, style = 'CommentsStyle')


#-----------------------------------CHANGE VARIABLE NAMES QUICKLY---------------------------------#
Friday = table.cell(0,4)
Fridayword = table.cell(3, 4)
WordDescription5 = table.cell(2,4)
Fridayevent = table.cell(1,4)
Friday.paragraphs[0].add_run("Friday"+str(FDate), style = 'Contentstyle').bold = True
WordDescription5.paragraphs[0].add_run("Word of the Day: "+ getword(2019,5,10), style = 'CommentsStyle')
Fridayword.paragraphs[0].add_run('Verminllion(n) - A vivid red to reddish orange color', style = 'CommentsStyle')
Fridayevent.paragraphs[0].add_run(Fevent, style = 'CommentsStyle')


document.save('testfile.docx')

