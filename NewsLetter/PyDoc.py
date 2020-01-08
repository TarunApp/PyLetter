from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt
import requests
import urllib
import urllib.request
from bs4 import BeautifulSoup
import string
from WordScrape import getword


# month = ""
# str(month)
# day = ""
# str(day)
# month = "december" #Wrap GUI Control or CLI
# day = "12"

# months = ["january", "february", "march", "april", "may", "june", "july", "august", "september", "october", "november", "december"]
# l = range(12+1)

# monthconversion = {}

# ki = 0
# for i in months:
# 	ki = ki + 1
# 	monthconversion[i] = (ki) 


def makedocument(month, day):
	url = "https://www.timeanddate.com/on-this-day" + "/" + month + "/" + day
	monthconversion = {'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6, 'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12}
	request = requests.get(url)		#<---------- Move this into a seperate function
	html_content = request.text

	soup = BeautifulSoup(html_content, "html.parser")

	#Find the tag and its attribute
	x = soup.find("h3", class_ ='otd-ttl')
	y = x.find_next("h3", class_ ='otd-ttl')
	g = y.find_next("h3", class_='otd-ttl')
	h = g.find_next("h3", class_ = 'otd-ttl')
	e = h.find_next("h3", class_ = 'otd-ttl')
	# print(type(x))

	def parsetotext(foundsoup):
		k = []
		for child in foundsoup:
			k.append(child)
		l = k[1]
		i = l.strip(' ')
		rcb = []
		for zx in k[0]:
			rcb.append(zx)
		return rcb[0]+"-"+i
	#------------------------------------------------------#
	#
	#8:14 - YESSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS
	                                     

	#Open Document
	document = Document()
	section = document.sections[0]

	#FORMATTING - LANDSCAPE
	new_width, new_height = section.page_height, section.page_width
	section.orientation = WD_ORIENTATION.LANDSCAPE 
	section.page_width = new_width
	section.page_height = new_height

	#Title styling
	style1 = document.styles['Normal']
	font1 = style1.font
	font1.name = 'Times New Roman'
	font1.size = Pt(30)

	#Styling
	header = section.header
	paragraphtitle = header.paragraphs[0]
	paragraphtitle.text = "The Verona Times"
	paragraphtitle.style = document.styles['Normal']
	paragraphtitle.paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


	#------------------STYLES------------------------------------------#

	#Add separate stylings for content types

	#HEADING
	obj_styles = document.styles
	obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
	obj_font = obj_charstyle.font
	obj_font.size = Pt(12)
	obj_font.name = 'Century Gothic'

	#CONTENT
	obj_charstyle2 = obj_styles.add_style('Contentstyle', WD_STYLE_TYPE.CHARACTER)
	obj_font2 = obj_charstyle2.font
	obj_font2.size = Pt(12)
	obj_font2.name = 'Century Gothic'

	#TITLE


	#Table Formats
	table = document.add_table(rows=5, cols=5)

	idate = int(day)
	Mdate = str(monthconversion[month]) + "/" + str(day) #" 11/5"
	Tdate =  str(monthconversion[month]) + "/" + str(idate+1) #" 11/6"
	WDate = str(monthconversion[month]) + "/" + str((idate + 2)) #" 11/7"
	ThDate =  str(monthconversion[month])+ "/" + str((idate + 3)) #" 11/8"
	FDate =  str(monthconversion[month]) + "/" + str((idate + 4))#" 11/9"

	Mevent = parsetotext(x)
	Tevent = parsetotext(y)
	Wevent = parsetotext(g)
	Thevent = parsetotext(h)
	Fevent = parsetotext(e)



	#  #---------------------------------------------------------------------------------------#
	#Monday
	Monday = table.cell(0,0)
	Mondayword = table.cell(3,0)
	WordDescription = table.cell(2,0)
	Mondayevent = table.cell(1,0)
	# Bolding
	Monday.paragraphs[0].add_run("Monday "+str(Mdate), style = 'Contentstyle').bold = True
	Mondayword.paragraphs[0].add_run( getword(2018, 12, 1), style = 'CommentsStyle') #Add getword here
	WordDescription.paragraphs[0].add_run("Word of the Day: ", style = 'CommentsStyle').bold = True
	Mondayevent.paragraphs[0].add_run( "On this day: " + Mevent, style = 'CommentsStyle')



	 #---------------------------------------------------------------------------------------#
	Tuesday = table.cell(0,1)
	Tuesdayword = table.cell(3, 1)
	WordDescription2 = table.cell(2,1)
	Tuesdayevent = table.cell(1,1)
	Tuesday.paragraphs[0].add_run("Tuesday "+str(Tdate), style = 'Contentstyle').bold = True
	WordDescription2.paragraphs[0].add_run("Word of the Day: ", style = 'CommentsStyle').bold = True
	Tuesdayword.paragraphs[0].add_run(getword(2018, 12, 2), style = 'CommentsStyle') #Add getword here
	Tuesdayevent.paragraphs[0].add_run( "On this day: " + Tevent, style = 'CommentsStyle')


	 #---------------------------------------------------------------------------------------#
	Wednesday = table.cell(0,2)
	Wednesdayword = table.cell(3, 2)
	WordDescription3 = table.cell(2,2)
	Wednesdayevent = table.cell(1,2)
	Wednesday.paragraphs[0].add_run("Wednesday "+str(WDate), style = 'Contentstyle').bold = True
	WordDescription3.paragraphs[0].add_run("Word of the Day: ", style = 'CommentsStyle').bold = True
	Wednesdayword.paragraphs[0].add_run(getword(2018, 12, 3), style = 'CommentsStyle') #Add getword here
	Wednesdayevent.paragraphs[0].add_run("On this day: " +Wevent, style = 'CommentsStyle')

	 
	 #---------------------------------------------------------------------------------------#
	Thursday= table.cell(0,3)
	Thursdayword = table.cell(3, 3)
	WordDescription4 = table.cell(2,3)
	Thursdayevent = table.cell(1,3)
	Thursday.paragraphs[0].add_run("Thursday "+str(ThDate), style = 'Contentstyle').bold = True
	WordDescription4.paragraphs[0].add_run("Word of the Day: ", style = 'CommentsStyle').bold = True
	Thursdayword.paragraphs[0].add_run(getword(2018, 12, 4), style = 'CommentsStyle') #Add getword here
	Thursdayevent.paragraphs[0].add_run("On this day: " +Thevent, style = 'CommentsStyle')


	#-----------------------------------CHANGE VARIABLE NAMES QUICKLY---------------------------------#
	Friday = table.cell(0,4)
	Fridayword = table.cell(3, 4)
	WordDescription5 = table.cell(2,4)
	Fridayevent = table.cell(1,4)
	Friday.paragraphs[0].add_run("Friday "+str(FDate), style = 'Contentstyle').bold = True
	WordDescription5.paragraphs[0].add_run("Word of the Day: ", style = 'CommentsStyle').bold = True
	Fridayword.paragraphs[0].add_run(getword(2018, 12, 5), style = 'CommentsStyle') #Add getword here
	Fridayevent.paragraphs[0].add_run("On this day: " +Fevent, style = 'CommentsStyle')


	document.save('{}.docx'.format(month + str(day)))

